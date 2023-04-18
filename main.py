import docx
import json
import os
import pandas as pd
import pypdfium2 as pdfium
import sys
import warnings

# Do not clog console with deprecation warnings
warnings.filterwarnings("ignore", category=DeprecationWarning)

from tqdm import tqdm
from os.path import join
from collections import Counter
from nltk import ngrams
from nltk.stem import WordNetLemmatizer
from nltk.stem.porter import PorterStemmer

# from nltk.corpus import stopwords

porter_stemmer = PorterStemmer()
wordnet_lemmatizer = WordNetLemmatizer()


[_, DOCUMENTS] = sys.argv
FILE_NAME_TO_SHORT_NAME = {
    "Motel": "motel",
    "Business Support Office": "bso",
    "Equity Development Office": "edo",
    "Resident Support Office": "rso",
    "Communication": "comm",
    "Admin": "admin",
}
MISC_FILE_EXT = [
    "xlsx",
    "html",
    # "pdf",
    "png",
    # "docx",
    "jpg",
    "jpeg",
    "zip",
    "pptx",
    "xls",
    "pages",
    # "csv",
    "m4a",
    # MAYBE GET TXT
    "txt",
    "conf",
    "mp4",
    "doc",
]


def get_ngrams(words):
    bi = [words[i : i + 2] for i in range(len(words) - 2 + 1)]
    tri = [words[i : i + 3] for i in range(len(words) - 3 + 1)]
    quad = [words[i : i + 4] for i in range(len(words) - 4 + 1)]
    return bi, tri, quad


def get_n_gram_counts(word_list, n_range=3, top_n=10):
    for nn in range(1, (n_range + 1)):
        print(" ")
        print(f"Processing {nn} ...")
        print(" ")
        ngram_counts = Counter(ngrams(word_list, nn))
        print(ngram_counts.most_common(top_n))


def find_cat_words(word_list, bi, tri, quad):
    with open(
        "C:/Users/ramosv/Desktop/MileHighHack/organize_chaos/Resources/lexicon.json",
        "r",
    ) as j:
        lex = json.loads(j.read())

    motel_words = get_key_gram_count(lex["Motel"], word_list, bi, tri, quad)
    bso_words = get_key_gram_count(
        lex["Business Support Office"], word_list, bi, tri, quad
    )
    edo_words = get_key_gram_count(
        lex["Equity Development Office"], word_list, bi, tri, quad
    )
    rso_issues_words = get_key_gram_count(
        lex["Resident Support Office"], word_list, bi, tri, quad
    )
    comm_words = get_key_gram_count(lex["Communication"], word_list, bi, tri, quad)
    admin_words = get_key_gram_count(lex["Admin"], word_list, bi, tri, quad)
    # thefax_words = get_key_gram_count(lex['The Fax'], word_list, bi, tri, quad)
    return motel_words, bso_words, edo_words, rso_issues_words, comm_words, admin_words


def get_key_gram_count(key_words, nvr_uni, nvr_bi, nvr_tri, nvr_quad):
    kw_unigrams = clean_keywords(key_words)
    kw_bigrams = [word.split() for word in kw_unigrams if len(word.split()) == 2]
    kw_trigrams = [word.split() for word in kw_unigrams if len(word.split()) == 3]
    kw_quadgrams = [word.split() for word in kw_unigrams if len(word.split()) == 4]
    uni_count = 0
    bi_count = 0
    tri_count = 0
    quad_count = 0
    for word in kw_unigrams:
        temp = nvr_uni.count(word)
        uni_count = uni_count + temp
    # print(f"There are a total of {uni_count} UNIGRAM MATCHES")
    for bigram in kw_bigrams:
        temp = nvr_bi.count(bigram) * 2
        bi_count = bi_count + temp
    # print(f"There are a total of {bi_count} BIGRAM MATCHES")
    for trigram in kw_trigrams:
        temp = nvr_tri.count(trigram) * 3
        tri_count = tri_count + temp
    # print(f"There are a total of {tri_count} TRIGRAM MATCHES")
    for quad in kw_quadgrams:
        temp = nvr_quad.count(quad) * 4
        quad_count = quad_count + temp
    # print(f"There are a total of {tri_count} TRIGRAM MATCHES")
    return uni_count + bi_count + tri_count + quad_count


def classify_on_model(
    output_path: str, unmatched_files: list[str], unmatched: list[str]
):
    """Classify the file's designation by the lexicon model."""
    files = [
        os.path.join(dp, f)
        for dp, dn, filenames in os.walk(os.getcwd())
        for f in filenames
        if f in unmatched_files
    ]

    df_list = []
    for file in tqdm(files):
        print(f"Processing {file} ...")
        # Do not move forward with anything that isn't docs or pdf
        if os.path.splitext(file)[1].lower() not in [".docx", ".pdf", ".txt"]:
            unmatched.append(file)
            continue

        txt = getText(file)

        w_list = tokenize(txt, [], type="word")
        bi, tri, quad = get_ngrams(w_list)
        (
            motel_words,
            bso_words,
            edo_words,
            rso_issues_words,
            comm_words,
            admin_words,
        ) = find_cat_words(w_list, bi, tri, quad)
        df_list.append(
            [
                file,
                len(w_list),
                motel_words,
                bso_words,
                edo_words,
                rso_issues_words,
                comm_words,
                admin_words,
            ]
        )

    df = pd.DataFrame(
        df_list,
        columns=[
            "Filename",
            "doc_word_count",
            "motel_words",
            "bso_words",
            "edo_words",
            "rso_issues_words",
            "comm_words",
            "admin_words",
        ],
    )
    # df['motel_words_perc'] = df['motel_words']/df['doc_word_count']
    # df['bso_words_perc'] = df['bso_words']/df['doc_word_count']
    # df['edo_words_perc'] = df['edo_words']/df['doc_word_count']
    # df['rso_issues_perc'] = df['rso_issues_words']/df['doc_word_count']
    # df['comm_words_perc'] = df['comm_words']/df['doc_word_count']
    # df['admin_perc'] = df['admin_words']/df['doc_word_count']
    # df['thefax_perc'] = df['thefax_words']/df['doc_word_count']

    print(df.shape)
    print(df)
    df.to_csv(output_path + "/RESULTS.csv")
    folderList = []
    for index, row in df.iterrows():
        # print(row['motel_words'],row['bso_words'],row['edo_words'], row['rso_issues_words'], row['comm_words'],row['admin_words'])
        isMax = max(
            row["motel_words"],
            row["bso_words"],
            row["edo_words"],
            row["rso_issues_words"],
            row["comm_words"],
            row["admin_words"],
        )

        if isMax == 0:
            folder = "misc"
        elif isMax == row["motel_words"]:
            folder = "motel"
        elif isMax == row["bso_words"]:
            folder = "bso"
        elif isMax == row["rso_issues_words"]:
            folder = "rso"
        elif isMax == row["comm_words"]:
            folder = "comm"
        elif isMax == row["admin_words"]:
            folder = "admin"
        elif isMax == row["edo_words"]:
            folder = "edo"
        else:
            raise Exception(
                "An unknown folder is being referenced as the main category. We cannot assign a file\n"
                + "to a folder we don't know about. Either include the folder in the above `if`-chain\n"
                + "or remove the lexicon matching to the unknown folder."
            )

        folderList.append([row["Filename"], folder])
    return folderList


def getText(filename: str):
    """Read the bytes of a file and return as string

    Returns (str): The content of the file, as string.
    """
    if os.path.splitext(filename)[1].lower() == ".docx":
        doc = docx.Document(filename)
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
        return "\n".join(fullText)
    elif os.path.splitext(filename)[1].lower() == ".pdf":
        return get_pdf_content(filename)
    elif os.path.splitext(filename)[1].lower() == ".txt":
        fullText: str
        with open(filename, "r") as txt_file:
            fullText = txt_file.read()
        txt_file.close()
        return fullText
    else:
        raise Exception(
            "Cannot parse the file because it is neither a `pdf`, `docx`, or `txt`."
        )


def filter_wordlist(word_list):
    return [word.lower() for word in word_list]


def clean_keywords(word_list):
    word_list = filter_wordlist(word_list)
    cleaned_keywords = stemming_lemming(word_list)
    return cleaned_keywords


def tokenize(text, exclude_words, type="word"):
    for spec_char in """!()-[]{};:'"\,<>/.?@#$%^&*_~""":
        text = text.replace(spec_char, "")
    exclude_words.extend(["also"])
    if type != "\n":
        text = text.replace("\n", "")
    if type == "word":
        lang_list = text.split()
        lang_list = filter_wordlist(lang_list)
        lang_list = stemming_lemming(lang_list)
        final_list = [word for word in lang_list if word not in exclude_words]
    elif type == "sentence":
        final_list = text.split(".")
    return final_list


def stemming_lemming(word_list, stem=True):
    if stem:
        words = [porter_stemmer.stem(word) for word in word_list]
    else:
        words = [wordnet_lemmatizer.lemmatize(word) for word in word_list]
    return words


def get_top_n_words(file_name, word_list, output, n=10):
    vocab_d = {}
    filtered_words = filter_wordlist(word_list)
    for word in filtered_words:
        if word not in vocab_d:
            vocab_d[word] = 1
        else:
            vocab_d[word] += 1

    df = pd.DataFrame(vocab_d.items(), columns=["Word", "Count"])
    df = df.sort_values(by="Count", ascending=False)
    df = df.head(n)
    print(df)
    df.to_csv(f"{output}/{file_name}_top_{n}_word_list.csv")


def get_pdf_content(file: str) -> str:
    """Derive the file contents of pdf"""
    try:
        pdf = pdfium.PdfDocument(file)
        contents = ""
        for i in range(len(pdf)):
            page = pdf[i]
            # Load a text page helper
            textpage = page.get_textpage()
            text_all = textpage.get_text_range()
            contents += text_all

        return contents
    except Exception as e:
        print(f"\n\nCannot parse PDF at '{file}'\n\n")
        raise e


def classify_on_file_name(file: str, matched: list, unmatched: list):
    """Classify the file's designation by name."""
    lexicon_dict: dict = read_lexicon_json()

    # Boolean to tell us whether we found a match
    found_match = False
    # Assign a dictionary to the file we are sorting
    sorted_file_dict = {"file": file}

    # For every category...
    for category in lexicon_dict.keys():
        # For every keyword in out lexicon array for the category...
        for keyword in lexicon_dict[category]:
            # If the keyword is in the file name, we have a match
            if keyword.lower() in file.lower():
                # If the category key isn't in the matches, then assign it
                if not "folder" in sorted_file_dict:
                    sorted_file_dict["folder"] = category
                    found_match = True
            # If no match, don't do shit
            else:
                continue

    if found_match:
        matched.append([file, FILE_NAME_TO_SHORT_NAME[sorted_file_dict["folder"]]])
    else:
        unmatched.append(file)


def classify_on_ext(file: str, matched: list, unmatched: list):
    """Classify the file's designation by extension."""
    if "." in file:
        extension = file.split(".").pop()
        if extension in MISC_FILE_EXT:
            matched.append([file, "misc"])
        else:
            unmatched.append(file)
    else:
        unmatched.append(file)


def read_lexicon_json() -> dict:
    """Read Lexicon JSON file."""
    with open(
        "C:/Users/ramosv/Desktop/MileHighHack/organize_chaos/Resources/lexicon.json",
        "r",
    ) as j:
        lex = json.loads(j.read())
    return lex


def organize_files(unorganized_documents_path: str, output_path: str):
    """Designate a file's placement give the lexicon matching. This function
    outputs a CSV with the absolute path to the file being tied to the folder
    name it should be moved to.

    Args:
        unorganized_documents_path (str): The absolute path to the unorganized `Documents` directory.
        output_path (str): The absolute path of the directory to output the cleaned up file structure, as CSV.
    """
    # Change directory into the unorganized documents directory
    os.chdir(unorganized_documents_path)

    # Grab all the file names
    # ---
    # NOTE: We leave out and knowledge about the directory name. The variable
    # below called `dir_names` would be the array of ALL folder names within the
    # `unorganized_documents_path` directory.
    files = [
        # `f` is the file name
        os.path.join(dir_path, f)
        for dir_path, dir_names, filenames in os.walk(os.getcwd())
        for f in filenames
    ]

    # Keep track of what was matched and what wasn't matched based on file name
    matched_files_on_file_name = []
    unmatched_files_on_file_name = []

    # Filter by file name
    for file in tqdm(files):
        classify_on_file_name(
            file, matched_files_on_file_name, unmatched_files_on_file_name
        )

    # Keep track of what was matched and what wasn't matched based on file extension
    matched_files_on_file_ext = []
    unmatched_files_on_file_ext = []
    # Filter by extension
    for file in tqdm(unmatched_files_on_file_name):
        classify_on_ext(files, matched_files_on_file_ext, unmatched_files_on_file_ext)

    # Keep track of what wasn't matched based on the model
    unmatched_files_on_model = []
    matched_files_on_model = classify_on_model(
        output_path, unmatched_files_on_file_ext, unmatched_files_on_model
    )

    # Aggregate results from all three steps (file name, file ext, and model)
    all_matched_files = (
        matched_files_on_file_name + matched_files_on_file_ext + matched_files_on_model
    )

    # Append the files that weren't matched from all the steps as miscellaneous
    for file in unmatched_files_on_model:
        all_matched_files.append([file, "misc"])

    # Put it all into a data frame
    df = pd.DataFrame(all_matched_files, columns=["File", "Folder"])
    #  Write it to the final output as CSV
    df.to_csv(join(output_path, "FinalOutput.csv"))


if "__main__" == __name__:
    output_path = "C:/Users/ramosv/Desktop/MileHighHack/organize_chaos/Output"
    unorganized_documents_path = DOCUMENTS
    organize_files(unorganized_documents_path, output_path)
